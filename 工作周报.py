# -*- coding:utf-8 -*-
"""
@Author: Felix
@File:   工作周报.py
@Desc:
@Create: 2020/09/22 11:22
"""
import pandas as pd
import matplotlib.pyplot as plot
from docx import Document
from docx.shared import Inches

# 生成的图标图片
images_path = "./images/PicExcel/plot1.jpg"


# 向Excel写入数据
def write_excel():
    df = pd.DataFrame({
        'id': [1, 2, 3],
        'name': ['zhang', 'li', 'wang1'],
        'age': [10, 20, 30]
    })

    df.set_index('id')
    print(df)
    # df.to_excel("./File/people.xlsx")
    print("Done")


# 读取Excel数据、生成图表
def read_Generate_execl():
    # 读取Excel
    student = pd.read_excel("./Excel/Student-1.xlsx")
    # sort_values 排序，inplace：True 对当前对象进行更改调整，ascending ：False反序
    student.sort_values(by="Score", inplace=True, ascending=False)
    # 用bar方法绘制柱状图
    plot.bar(student.Name, student.Score, color="blue")
    # 设置标题，x轴，y轴
    plot.title("student", fontsize=20)
    plot.xlabel("Name")
    plot.ylabel("Score")

    # 如果x轴字数太多，那么则利用xticks方法rotation传参设置旋转角度
    plot.xticks(student.Name, rotation="90")
    # 布局方式：紧凑布局
    plot.tight_layout()
    # # 最终呈现效果
    # plot.show()
    # 保存数据至图片文件
    plot.savefig(images_path)
    print("生成成功")
    return student


# 操作Word
def operate_word(student):
    # 分数第一的学生
    first_student = student.iloc[0]
    # # 创建一个新的空文档
    # document = Document()
    # 读取一个存在的文档
    document = Document("./Word/w1.docx")
    # print(help(document))
    # 增加标题
    document.add_heading('数据分析报告', level=0)
    # 增加段落
    para1 = document.add_paragraph('分数第一的学生是：')
    # 延长段落  .bold = True  加粗字体
    para1.add_run(str(first_student["Name"])).bold = True
    para1.add_run("，分数为：")
    para1.add_run(str(first_student["Score"])).bold = True
    # 在段落之前增加一个新段落
    para2 = para1.insert_paragraph_before(f'总共有{len(student.Name)}名学生参加了考试，学生考试总体情况：')
    # # 增加分页符
    # document.add_page_break()
    # 增加一个表格
    table = document.add_table(rows=len(student.Name) + 1, cols=2)
    # 表格样式
    table.style = "Light Shading Accent 1"
    table.cell(0, 0).text = "学生姓名"
    table.cell(0, 1).text = "学生分数"
    # 循环插入数据
    # print(enumerate(student.iterrows()))
    for i, (index, row) in enumerate(student.iterrows()):
        # print(f"i: {i} index: {index} row: {row}")
        table.cell(i + 1, 0).text = str(row["Name"])
        table.cell(i + 1, 1).text = str(row["Score"])
    # # 第一行第三列
    # cell_1_3 = table.cell(0, 2)
    # # 赋值
    # cell_1_3.text = "第一行第三列"
    # # 第二行
    # row_2 = table.rows[1]
    # row_2.cells[0].text = "第二行第一列"
    # row_2.cells[1].text = "第二行第二列"
    # 增加图片
    document.add_picture(images_path)
    # 保存文档到指定文件
    document.save("./Word/w1.docx")
    print("Done")


if __name__ == '__main__':
    student = read_Generate_execl()
    operate_word(student)
